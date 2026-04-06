"""Word 문서 생성 엔진 — python-docx 래퍼"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import structlog

logger = structlog.get_logger(__name__)


class WordEngine:
    """python-docx 기반 Word 생성 엔진"""

    def create(
        self,
        title: str,
        content_blocks: list[dict],
        output_path: Path,
    ) -> dict:
        """Word 파일 생성

        Args:
            title: 문서 제목
            content_blocks: 콘텐츠 블록 목록
                            [{"type": "heading|paragraph|list", "text": "...", "level": 1}]
            output_path: 저장 경로

        Returns:
            dict: 생성 결과 정보
        """
        doc = Document()
        
        # ── 문서 제목 설정 ──
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ── 콘텐츠 블록 작성 ──
        for block in content_blocks:
            b_type = block.get("type", "paragraph")
            text = block.get("text", "")
            
            if b_type == "heading":
                level = block.get("level", 1)
                doc.add_heading(text, level=min(level, 9))
            elif b_type == "list":
                doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph(text)

        doc.save(str(output_path))
        file_size = output_path.stat().st_size

        logger.info(
            "word_created",
            blocks=len(content_blocks),
            size_bytes=file_size,
            path=str(output_path),
        )

        return {
            "block_count": len(content_blocks),
            "size_bytes": file_size,
        }
